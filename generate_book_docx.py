"""
Generate DOCX files for books
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

def generate_book_docx(book_data):
    """Generate a DOCX file for a book"""
    doc = Document()
    
    # Set up default font
    style = doc.styles['Normal']
    font = style.font
    font.name = book_data.get('font_family', 'Times New Roman')
    font.size = Pt(book_data.get('font_size', 11))
    
    # Title page
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(book_data.get('book_title', 'Untitled'))
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    
    doc.add_paragraph()  # Spacer
    
    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_run = author.add_run(f"by {book_data.get('author_name', 'Author')}")
    author_run.font.size = Pt(16)
    
    if book_data.get('genre'):
        doc.add_paragraph()
        genre = doc.add_paragraph()
        genre.alignment = WD_ALIGN_PARAGRAPH.CENTER
        genre_run = genre.add_run(book_data['genre'])
        genre_run.font.size = Pt(14)
        genre_run.font.italic = True
    
    doc.add_page_break()
    
    # Dedication
    if book_data.get('dedication'):
        doc.add_paragraph()
        doc.add_paragraph()
        dedication = doc.add_paragraph()
        dedication.alignment = WD_ALIGN_PARAGRAPH.CENTER
        dedication_run = dedication.add_run(book_data['dedication'])
        dedication_run.font.italic = True
        doc.add_page_break()
    
    # About the Author
    if book_data.get('about_author'):
        about_title = doc.add_heading('About the Author', level=2)
        doc.add_paragraph(book_data['about_author'])
        doc.add_page_break()
    
    # Table of Contents
    toc_title = doc.add_heading('Table of Contents', level=1)
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Calculate chapter starting pages
    current_page = 1  # title page
    if book_data.get('dedication'):
        current_page += 1
    if book_data.get('about_author'):
        current_page += 1
    current_page += 1  # TOC page itself
    
    for i, chapter in enumerate(book_data.get('chapters', [])):
        chapter_title = chapter.get('title', f'Chapter {chapter.get("number", i+1)}')
        toc_entry = doc.add_paragraph()
        toc_entry.add_run(f"{chapter_title} ")
        toc_entry.add_run('.' * 50)
        toc_entry.add_run(f" {current_page}")
        
        # Estimate pages for this chapter
        content = chapter.get('content', '')
        word_count = len(content.split())
        estimated_pages = max(1, (word_count + 350) // 400)
        current_page += estimated_pages
    
    doc.add_page_break()
    
    # Chapters
    for chapter in book_data.get('chapters', []):
        # Chapter number
        chapter_num = doc.add_paragraph()
        chapter_num.alignment = WD_ALIGN_PARAGRAPH.LEFT
        num_run = chapter_num.add_run(str(chapter.get('number', '')))
        num_run.font.size = Pt(18)
        num_run.font.bold = True
        
        # Chapter title
        chapter_title = doc.add_heading(chapter.get('title', ''), level=1)
        
        doc.add_paragraph()  # Spacer
        
        # Chapter content
        content = chapter.get('content', '')
        paragraphs = content.split('\n\n')
        for para in paragraphs:
            if para.strip():
                p = doc.add_paragraph(para.strip())
                p.paragraph_format.first_line_indent = Inches(0.3)
                p.paragraph_format.line_spacing = 1.5
        
        doc.add_page_break()
    
    # Save to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer.getvalue()
